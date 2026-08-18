# # """
# # document_export.py
# # ────────────────────────────────────────────────────────────────
# # Post-processing utility that turns an LLM's markdown response into a
# # downloadable .docx and uploads it to Google Drive (service account).

# # IMPORTANT: This module is 100% decoupled from the RAG/LLM chain in
# # api.py. It only ever receives the *final* markdown string the LLM
# # already produced — it never sees the prompt, the retriever, or the
# # chain internals. Nothing here changes model behaviour.

# # SETUP REQUIRED:
# #   1. Create a Google Cloud project → enable the "Google Drive API".
# #   2. Create a Service Account → generate a JSON key.
# #   3. Save the key file somewhere on the server, e.g. ./service_account.json
# #   4. Set env var: GOOGLE_SERVICE_ACCOUNT_FILE=/path/to/service_account.json
# #   5. pip install google-api-python-client google-auth
# # """

# # import os
# # import re
# # from dataclasses import dataclass

# # from docx import Document as DocxDocument
# # from docx.enum.text import WD_ALIGN_PARAGRAPH

# # from googleapiclient.discovery import build
# # from googleapiclient.http import MediaFileUpload
# # from google.oauth2 import service_account


# # # ══════════════════════════════════════════════════════════════
# # # 1. CONFIG
# # # ══════════════════════════════════════════════════════════════

# # SCOPES = ["https://www.googleapis.com/auth/drive.file"]  # only files this account creates
# # EXPORT_DIR = os.getenv("DOC_EXPORT_DIR", "generated_docs")

# # os.makedirs(EXPORT_DIR, exist_ok=True)


# # @dataclass
# # class DocumentResult:
# #     title: str
# #     doc_type: str
# #     download_url: str
# #     view_url: str
# #     file_id: str


# # # ══════════════════════════════════════════════════════════════
# # # 2. INTENT DETECTION — keyword heuristic on the USER QUERY only
# # #    (never inspects/changes LLM output logic)
# # # ══════════════════════════════════════════════════════════════

# # DOCUMENT_TRIGGERS = [
# #     "draft", "write a judgment", "write the judgment", "prepare a judgment",
# #     "generate a document", "generate document", "download",
# #     "export", "make up a document", "create a document", "judgment for",
# #     "decree", "write up the judgment", "full judgment", "complete judgment",
# #     "write out the order", "prepare the order",
# # ]


# # def should_generate_document(query: str) -> bool:
# #     """Cheap, dependency-free heuristic. Tune DOCUMENT_TRIGGERS to taste,
# #     or swap this for an LLM-classifier call later without touching
# #     anything else in this module."""
# #     q = query.lower()
# #     return any(trigger in q for trigger in DOCUMENT_TRIGGERS)


# # # ══════════════════════════════════════════════════════════════
# # # 3. MARKDOWN → DOCX
# # #    Preserves the LLM's exact wording — this only maps markdown
# # #    syntax to Word formatting, it does not rewrite any content.
# # # ══════════════════════════════════════════════════════════════

# # _INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")


# # def _add_inline_runs(paragraph, text: str) -> None:
# #     """Parse **bold**, *italic*, ***bold italic*** inline spans."""
# #     pos = 0
# #     for match in _INLINE_RE.finditer(text):
# #         if match.start() > pos:
# #             paragraph.add_run(text[pos:match.start()])
# #         token = match.group(0)
# #         if token.startswith("***") and token.endswith("***"):
# #             run = paragraph.add_run(token[3:-3])
# #             run.bold = True
# #             run.italic = True
# #         elif token.startswith("**") and token.endswith("**"):
# #             run = paragraph.add_run(token[2:-2])
# #             run.bold = True
# #         else:
# #             run = paragraph.add_run(token[1:-1])
# #             run.italic = True
# #         pos = match.end()
# #     if pos < len(text):
# #         paragraph.add_run(text[pos:])


# # def markdown_to_docx(markdown_text: str, title: str) -> str:
# #     """Convert markdown → .docx, preserving headings, bold/italic,
# #     bullet & numbered lists, and paragraph breaks. Returns local filepath."""
# #     doc = DocxDocument()

# #     heading = doc.add_heading(title, level=0)
# #     heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# #     for raw_line in markdown_text.splitlines():
# #         line = raw_line.rstrip()

# #         if not line.strip():
# #             continue

# #         heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
# #         if heading_match:
# #             level = min(len(heading_match.group(1)), 4)
# #             doc.add_heading(heading_match.group(2).strip(), level=level)
# #             continue

# #         if re.match(r"^(-{3,}|_{3,}|\*{3,})\s*$", line.strip()):
# #             doc.add_paragraph("─" * 60)
# #             continue

# #         if re.match(r"^\s*[-*]\s+", line):
# #             p = doc.add_paragraph(style="List Bullet")
# #             _add_inline_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
# #             continue

# #         if re.match(r"^\s*\d+[.)]\s+", line):
# #             p = doc.add_paragraph(style="List Number")
# #             _add_inline_runs(p, re.sub(r"^\s*\d+[.)]\s+", "", line))
# #             continue

# #         p = doc.add_paragraph()
# #         _add_inline_runs(p, line)

# #     safe_title = re.sub(r"[^\w\-]+", "_", title).strip("_")[:80] or "Document"
# #     filepath = os.path.join(EXPORT_DIR, f"{safe_title}.docx")
# #     doc.save(filepath)
# #     return filepath


# # # ══════════════════════════════════════════════════════════════
# # # 4. GOOGLE DRIVE UPLOAD (service account)
# # # ══════════════════════════════════════════════════════════════

# # _drive_service = None


# # def _get_drive_service():
# #     global _drive_service
# #     if _drive_service is None:
# #         service_account_file = os.getenv("GOOGLE_SERVICE_ACCOUNT_FILE", "service_account.json")
# #         if not os.path.exists(service_account_file):
# #             raise RuntimeError(
# #                 f"Service account key not found at '{service_account_file}'. "
# #                 "Set GOOGLE_SERVICE_ACCOUNT_FILE (in .env) or place the JSON key there."
# #             )
# #         creds = service_account.Credentials.from_service_account_file(
# #             service_account_file, scopes=SCOPES
# #         )
# #         _drive_service = build("drive", "v3", credentials=creds)
# #     return _drive_service


# # def upload_docx_to_drive(filepath: str, filename: str) -> DocumentResult:
# #     """Upload local .docx to Drive, make it link-shareable, return URLs."""
# #     service = _get_drive_service()

# #     file_metadata = {
# #         "name": filename,
# #         "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
# #     }
# #     media = MediaFileUpload(
# #         filepath,
# #         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
# #         resumable=False,
# #     )

# #     uploaded = service.files().create(
# #         body=file_metadata, media_body=media, fields="id"
# #     ).execute()
# #     file_id = uploaded["id"]

# #     service.permissions().create(
# #         fileId=file_id,
# #         body={"role": "reader", "type": "anyone"},
# #         fields="id",
# #     ).execute()

# #     meta = service.files().get(
# #         fileId=file_id, fields="webContentLink,webViewLink,name"
# #     ).execute()

# #     return DocumentResult(
# #         title=filename,
# #         doc_type="DOCX",
# #         download_url=meta["webContentLink"],
# #         view_url=meta["webViewLink"],
# #         file_id=file_id,
# #     )


# # # ══════════════════════════════════════════════════════════════
# # # 5. ONE-CALL ENTRYPOINT for api.py to use
# # # ══════════════════════════════════════════════════════════════

# # def export_markdown_as_docx(markdown_text: str, title: str) -> DocumentResult:
# #     """markdown → docx → Drive upload, in a single call."""
# #     filepath = markdown_to_docx(markdown_text, title)
# #     try:
# #         return upload_docx_to_drive(filepath, os.path.basename(filepath))
# #     finally:
# #         if os.getenv("DOC_EXPORT_CLEANUP", "true").lower() == "true":
# #             try:
# #                 os.remove(filepath)
# #             except OSError:
# #                 pass


















# """
# document_export.py
# ────────────────────────────────────────────────────────────────
# Post-processing utility that turns an LLM's markdown response into a
# downloadable .docx file, saved locally and served by our own FastAPI
# endpoint (see /documents/{filename} in api.py).

# IMPORTANT: This module is 100% decoupled from the RAG/LLM chain in
# api.py. It only ever receives the *final* markdown string the LLM
# already produced — it never sees the prompt, the retriever, or the
# chain internals. Nothing here changes model behaviour.

# NOTE ON GOOGLE DRIVE: an earlier version of this module uploaded to
# Google Drive via a service account. That doesn't work on a personal
# (non-Workspace) Google account — service accounts have no storage
# quota of their own outside Shared Drives / domain-wide delegation,
# both of which require a paid Google Workspace org. So instead, we
# just save the .docx locally and serve it directly — simpler, no
# Google credentials needed at all, and no quota limits.
# """

# import os
# import re
# from dataclasses import dataclass

# from docx import Document as DocxDocument
# from docx.enum.text import WD_ALIGN_PARAGRAPH


# # ══════════════════════════════════════════════════════════════
# # 1. CONFIG
# # ══════════════════════════════════════════════════════════════

# EXPORT_DIR = os.getenv("DOC_EXPORT_DIR", "generated_docs")
# # Base URL your server is reachable at. Update this when you deploy
# # (e.g. https://api.yourdomain.com) — for local dev, localhost is fine.
# PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://127.0.0.1:8000")

# os.makedirs(EXPORT_DIR, exist_ok=True)


# @dataclass
# class DocumentResult:
#     title: str
#     doc_type: str
#     download_url: str
#     filename: str


# # ══════════════════════════════════════════════════════════════
# # 2. INTENT DETECTION — keyword heuristic on the USER QUERY only
# #    (never inspects/changes LLM output logic)
# # ══════════════════════════════════════════════════════════════

# DOCUMENT_TRIGGERS = [
#     "draft", "write a judgment", "write the judgment", "prepare a judgment",
#     "generate a document", "generate document", "download",
#     "export", "make up a document", "create a document", "judgment for",
#     "decree", "write up the judgment", "full judgment", "complete judgment",
#     "write out the order", "prepare the order",
# ]


# def should_generate_document(query: str) -> bool:
#     """Cheap, dependency-free heuristic. Tune DOCUMENT_TRIGGERS to taste,
#     or swap this for an LLM-classifier call later without touching
#     anything else in this module."""
#     q = query.lower()
#     return any(trigger in q for trigger in DOCUMENT_TRIGGERS)


# # ══════════════════════════════════════════════════════════════
# # 3. MARKDOWN → DOCX
# #    Preserves the LLM's exact wording — this only maps markdown
# #    syntax to Word formatting, it does not rewrite any content.
# # ══════════════════════════════════════════════════════════════

# _INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")


# def _add_inline_runs(paragraph, text: str) -> None:
#     """Parse **bold**, *italic*, ***bold italic*** inline spans."""
#     pos = 0
#     for match in _INLINE_RE.finditer(text):
#         if match.start() > pos:
#             paragraph.add_run(text[pos:match.start()])
#         token = match.group(0)
#         if token.startswith("***") and token.endswith("***"):
#             run = paragraph.add_run(token[3:-3])
#             run.bold = True
#             run.italic = True
#         elif token.startswith("**") and token.endswith("**"):
#             run = paragraph.add_run(token[2:-2])
#             run.bold = True
#         else:
#             run = paragraph.add_run(token[1:-1])
#             run.italic = True
#         pos = match.end()
#     if pos < len(text):
#         paragraph.add_run(text[pos:])


# def markdown_to_docx(markdown_text: str, title: str) -> str:
#     """Convert markdown → .docx, preserving headings, bold/italic,
#     bullet & numbered lists, and paragraph breaks. Returns local filepath."""
#     doc = DocxDocument()

#     heading = doc.add_heading(title, level=0)
#     heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     for raw_line in markdown_text.splitlines():
#         line = raw_line.rstrip()

#         if not line.strip():
#             continue

#         heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
#         if heading_match:
#             level = min(len(heading_match.group(1)), 4)
#             doc.add_heading(heading_match.group(2).strip(), level=level)
#             continue

#         if re.match(r"^(-{3,}|_{3,}|\*{3,})\s*$", line.strip()):
#             doc.add_paragraph("─" * 60)
#             continue

#         if re.match(r"^\s*[-*]\s+", line):
#             p = doc.add_paragraph(style="List Bullet")
#             _add_inline_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
#             continue

#         if re.match(r"^\s*\d+[.)]\s+", line):
#             p = doc.add_paragraph(style="List Number")
#             _add_inline_runs(p, re.sub(r"^\s*\d+[.)]\s+", "", line))
#             continue

#         p = doc.add_paragraph()
#         _add_inline_runs(p, line)

#     safe_title = re.sub(r"[^\w\-]+", "_", title).strip("_")[:80] or "Document"
#     filepath = os.path.join(EXPORT_DIR, f"{safe_title}.docx")
#     doc.save(filepath)
#     return filepath


# # ══════════════════════════════════════════════════════════════
# # 4. ONE-CALL ENTRYPOINT for api.py to use
# # ══════════════════════════════════════════════════════════════

# def export_markdown_as_docx(markdown_text: str, title: str) -> DocumentResult:
#     """
#     markdown → docx, saved in EXPORT_DIR, with a download URL pointing
#     at our own /documents/{filename} endpoint (see api.py).

#     Note: files persist in EXPORT_DIR (they're not deleted after this
#     call, since the download endpoint needs to serve them later). Set
#     up your own retention/cleanup job if this directory needs to stay
#     small over time.
#     """
#     filepath = markdown_to_docx(markdown_text, title)
#     filename = os.path.basename(filepath)
#     return DocumentResult(
#         title=title,
#         doc_type="DOCX",
#         download_url=f"{PUBLIC_BASE_URL}/documents/{filename}",
#         filename=filename,
#     )
















"""
document_export.py
────────────────────────────────────────────────────────────────
Post-processing utility that turns an LLM's markdown response into a
downloadable .docx file, saved locally and served by our own FastAPI
endpoint (see /documents/{filename} in api.py).

IMPORTANT: This module is 100% decoupled from the RAG/LLM chain in
api.py. It only ever receives the *final* markdown string the LLM
already produced — it never sees the prompt, the retriever, or the
chain internals. Nothing here changes model behaviour.

NOTE ON GOOGLE DRIVE: an earlier version of this module uploaded to
Google Drive via a service account. That doesn't work on a personal
(non-Workspace) Google account — service accounts have no storage
quota of their own outside Shared Drives / domain-wide delegation,
both of which require a paid Google Workspace org. So instead, we
just save the .docx locally and serve it directly — simpler, no
Google credentials needed at all, and no quota limits.
"""

import os
import re
from dataclasses import dataclass

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ══════════════════════════════════════════════════════════════
# 1. CONFIG
# ══════════════════════════════════════════════════════════════

EXPORT_DIR = os.getenv("DOC_EXPORT_DIR", "generated_docs")
# Base URL your server is reachable at. Update this when you deploy
# (e.g. https://api.yourdomain.com) — for local dev, localhost is fine.
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://127.0.0.1:8000")

os.makedirs(EXPORT_DIR, exist_ok=True)


@dataclass
class DocumentResult:
    title: str
    doc_type: str
    download_url: str
    filename: str


# ══════════════════════════════════════════════════════════════
# 2. INTENT DETECTION — keyword heuristic on the USER QUERY only
#    (never inspects/changes LLM output logic)
# ══════════════════════════════════════════════════════════════

DOCUMENT_NOUNS = [
    "document", "docx", "doc file", "docs file", "word file", "word document",
    "pdf", "downloadable file", ".doc", ".docx", "file for download",
]

ACTION_VERBS = [
    "create", "generate", "make", "give me", "prepare", "download",
    "export", "provide", "attach", "produce", "share",
]


def should_generate_document(query: str) -> bool:
    """
    Only trigger when the query BOTH names a file/document type AND uses
    an action verb requesting it — e.g. "give me a docx", "download the
    document", "export this as a word file".

    Plain judgment-drafting requests ("draft a judgment for...", "write
    the decree") do NOT trigger this on their own, even though that's
    the tool's core function — otherwise nearly every query would match.
    """
    q = query.lower()
    has_doc_noun = any(noun in q for noun in DOCUMENT_NOUNS)
    has_action = any(verb in q for verb in ACTION_VERBS)
    return has_doc_noun and has_action


# ══════════════════════════════════════════════════════════════
# 3. MARKDOWN → DOCX
#    Preserves the LLM's exact wording — this only maps markdown
#    syntax to Word formatting, it does not rewrite any content.
# ══════════════════════════════════════════════════════════════

_INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Parse **bold**, *italic*, ***bold italic*** inline spans."""
    pos = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("***") and token.endswith("***"):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(markdown_text: str, title: str) -> str:
    """Convert markdown → .docx, preserving headings, bold/italic,
    bullet & numbered lists, and paragraph breaks. Returns local filepath."""
    doc = DocxDocument()

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        if not line.strip():
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            continue

        if re.match(r"^(-{3,}|_{3,}|\*{3,})\s*$", line.strip()):
            doc.add_paragraph("─" * 60)
            continue

        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
            continue

        if re.match(r"^\s*\d+[.)]\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\s*\d+[.)]\s+", "", line))
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, line)

    safe_title = re.sub(r"[^\w\-]+", "_", title).strip("_")[:80] or "Document"
    filepath = os.path.join(EXPORT_DIR, f"{safe_title}.docx")
    doc.save(filepath)
    return filepath


# ══════════════════════════════════════════════════════════════
# 4. ONE-CALL ENTRYPOINT for api.py to use
# ══════════════════════════════════════════════════════════════

def export_markdown_as_docx(markdown_text: str, title: str) -> DocumentResult:
    """
    markdown → docx, saved in EXPORT_DIR, with a download URL pointing
    at our own /documents/{filename} endpoint (see api.py).

    Note: files persist in EXPORT_DIR (they're not deleted after this
    call, since the download endpoint needs to serve them later). Set
    up your own retention/cleanup job if this directory needs to stay
    small over time.
    """
    filepath = markdown_to_docx(markdown_text, title)
    filename = os.path.basename(filepath)
    return DocumentResult(
        title=title,
        doc_type="DOCX",
        download_url=f"{PUBLIC_BASE_URL}/documents/{filename}",
        filename=filename,
    )